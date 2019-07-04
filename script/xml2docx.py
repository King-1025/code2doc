#!/usr/bin/env python3

# -*- coding: UTF-8; -*-

try:
  from xml.dom.minidom import parse
  from docx import Document
  from docx.shared import Pt as pt
except ModuleNotFoundError as e:
  print(e)

import sys

def from_xml_to_create_docx(input_xml,output_docx):
   with open(input_xml,"r") as fxml:
      tree = parse(fxml)
      root = tree.documentElement
      part_list = root.getElementsByTagName("part")
      doc = None
      count = 0
      ignore_count = 0
      for part in part_list:
        ignore=part.getAttribute("ignore")
        if ignore == "yes":
           ignore_count+=1
           continue
        if doc == None:
           doc = Document()

        pgh = doc.add_paragraph()
        pgh.paragraph_format.line_spacing = pt(12)
        header = part.getElementsByTagName("header")[0]
        run=pgh.add_run(header.childNodes[0].data)
        font_size=header.getAttribute("font-size")
        if font_size != None:
           run.font.size=pt(float(font_size))
        bold=header.getAttribute("bold")
        if bold == "yes":
           run.bold=True
#       print(bold)
        font_name=header.getAttribute("font-name")
        if font_name != None:
           run.font.name=font_name

#       pgh = doc.add_paragraph()
        pgh.paragraph_format.line_spacing = pt(12)
        content = part.getElementsByTagName("content")[0]
        run=pgh.add_run(content.childNodes[0].data)
        font_size=content.getAttribute("font-size")
        if font_size != None:
           run.font.size=pt(float(font_size))
        font_name=content.getAttribute("font-name")
        if font_name != None:
           run.font.name=font_name
        
        print("+-------header--------+")
        print("%s" % header.childNodes[0].data)
        print("+---------------------+")
#        print("--------content--------")
#        print("%s" % content.childNodes[0].data)
        print("finished part --> "+str(count))
        count+=1
        print("")
      if doc != None:
         doc.save(output_docx)
         print("Success! total:"+str(count+ignore_count)+" handle:"+str(count)+" ignore:"+str(ignore_count)+" output:"+output_docx)
      else:
         print("Not found any parts in %s" % input_xml)

def app(input_xml,output_docx="out.docx"):
    from_xml_to_create_docx(input_xml,output_docx)

if __name__ == "__main__":
   argc = len(sys.argv)
   if argc == 2 or argc == 3:
     if argc == 2:
        app(str(sys.argv[1]))
   elif argc == 3:
        app(str(sys.argv[1]),str(sys.argv[2]))
   else:
     print("Usage:%s <xml_file> [out_docx]" % str(sys.argv[0]))
